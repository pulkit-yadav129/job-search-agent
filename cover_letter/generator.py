"""
generator.py
Generates a personalised, ATS-optimised cover letter for a specific job.
"""

import re
from langchain_openai import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain

from agent.prompts import COVER_LETTER_PROMPT


def generate_cover_letter(
    profile: dict,
    job_title: str,
    company: str,
    location: str,
    job_description: str,
    llm: ChatOpenAI,
    tone: str = "professional"
) -> str:
    """
    Generate a tailored cover letter.

    Args:
        profile:         Structured user profile dict.
        job_title:       Target job title.
        company:         Company name.
        location:        Job location string.
        job_description: Full or partial JD text.
        llm:             ChatOpenAI instance.
        tone:            'professional' | 'enthusiastic' | 'concise'

    Returns:
        Cover letter as a plain text string.
    """
    tone_instruction = {
        "professional": "Use a formal, polished tone.",
        "enthusiastic": "Use an energetic, passionate tone that shows genuine excitement.",
        "concise":      "Keep it under 250 words — clear and punchy.",
    }.get(tone, "Use a professional tone.")

    # Inject tone into prompt dynamically
    prompt_text = COVER_LETTER_PROMPT + f"\n\nTone instruction: {tone_instruction}"

    prompt = PromptTemplate(
        input_variables=["profile", "job_title", "company", "location", "job_description"],
        template=prompt_text
    )
    chain = LLMChain(llm=llm, prompt=prompt)

    letter = chain.run(
        profile=str(profile),
        job_title=job_title,
        company=company,
        location=location,
        job_description=job_description[:3000]
    ).strip()

    return letter


def format_cover_letter_docx(letter_text: str, candidate_name: str = "Candidate") -> bytes:
    """
    Package the cover letter as a .docx file (in-memory bytes).
    Returns raw bytes for Streamlit download.
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from io import BytesIO

    doc = Document()

    # Margins
    section = doc.sections[0]
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Header line
    header = doc.add_paragraph()
    run = header.add_run(candidate_name)
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()  # spacer

    # Body — split into paragraphs
    for para in letter_text.split("\n\n"):
        para = para.strip()
        if para:
            p = doc.add_paragraph(para)
            p.paragraph_format.space_after = Pt(10)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
